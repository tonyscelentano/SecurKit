"""Real-parser validation: build a docx/xlsx with the canonical Python libraries,
run it through scrub_office, then reopen with the same libraries and confirm
both that they open without errors AND that the document body survives.

If python-docx or openpyxl rejects the scrubbed file, the scrubber broke the
schema — no matter how clean the XML looked. This is the test that catches
'lxml round-trip looked fine but Office can't parse it'.
"""

from __future__ import annotations

import io
import zipfile

import openpyxl
from docx import Document

from securkit.scrubber import scrub_office


def _all_decompressed(data: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return b"".join(z.read(name) for name in z.namelist())


def test_python_docx_can_open_scrubbed_document() -> None:
    """Build a real docx with python-docx, scrub it, reopen — must not raise."""
    doc = Document()
    doc.core_properties.author = "REAL_AUTHOR_NAME"
    doc.core_properties.last_modified_by = "REAL_LAST_MODIFIED"
    doc.core_properties.title = "Confidential Memo"
    doc.add_heading("Q3 Compliance Findings", level=1)
    doc.add_paragraph("This is the cover summary paragraph.")
    doc.add_paragraph("A second paragraph with more body text.")
    doc.add_heading("Detail", level=2)
    doc.add_paragraph("Sensitive evidence goes here in the real workflow.")

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    assert b"REAL_AUTHOR_NAME" in _all_decompressed(raw)

    clean_bytes, findings = scrub_office(raw)

    # 1. The author bytes are gone from the package
    all_content = _all_decompressed(clean_bytes)
    assert b"REAL_AUTHOR_NAME" not in all_content
    assert b"REAL_LAST_MODIFIED" not in all_content
    assert b"Confidential Memo" not in all_content

    # 2. python-docx can REOPEN the scrubbed file without errors. This is the
    #    proof that we didn't break the schema while scrubbing.
    reopened = Document(io.BytesIO(clean_bytes))

    # 3. The document body survives the round-trip
    paras = [p.text for p in reopened.paragraphs]
    assert "Q3 Compliance Findings" in paras
    assert "This is the cover summary paragraph." in paras
    assert "Sensitive evidence goes here in the real workflow." in paras

    # 4. Metadata properties on the reopened doc are blank
    assert not (reopened.core_properties.author or "").strip()
    assert not (reopened.core_properties.last_modified_by or "").strip()
    assert not (reopened.core_properties.title or "").strip()


def test_openpyxl_can_open_scrubbed_workbook() -> None:
    """Same proof for xlsx: openpyxl must reopen the scrubbed workbook."""
    wb = openpyxl.Workbook()
    wb.properties.creator = "REAL_AUTHOR_NAME"
    wb.properties.lastModifiedBy = "REAL_LAST_MODIFIED"
    wb.properties.title = "Confidential Spreadsheet"
    ws = wb.active
    ws.title = "Findings"
    ws["A1"] = "Account"
    ws["B1"] = "Amount"
    ws["A2"] = "Operating"
    ws["B2"] = 12345.67
    ws["A3"] = "Reserve"
    ws["B3"] = 9876.54

    buf = io.BytesIO()
    wb.save(buf)
    raw = buf.getvalue()
    assert b"REAL_AUTHOR_NAME" in _all_decompressed(raw)

    clean_bytes, _ = scrub_office(raw)

    # No leaked author bytes anywhere in the package
    all_content = _all_decompressed(clean_bytes)
    assert b"REAL_AUTHOR_NAME" not in all_content
    assert b"REAL_LAST_MODIFIED" not in all_content

    # openpyxl can REOPEN — schema is intact (this is the real check)
    reopened = openpyxl.load_workbook(io.BytesIO(clean_bytes))
    ws2 = reopened["Findings"]
    assert ws2["A1"].value == "Account"
    assert ws2["B2"].value == 12345.67
    assert ws2["A3"].value == "Reserve"
    # NOTE: not asserting openpyxl.properties.creator / lastModifiedBy here —
    # openpyxl's accessor returns its own library defaults ('openpyxl') when
    # the underlying XML is empty, which gives false positives. The byte check
    # above on the decompressed package is the authoritative scrub verification.


def test_customxml_variable_indices_stubbed() -> None:
    """customXml/item*.xml at non-1 indices must be stubbed (not deleted —
    that would break [Content_Types].xml manifest references)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED) as z:
        z.writestr("[Content_Types].xml", '<?xml version="1.0"?><Types xmlns="x"/>')
        z.writestr("docProps/core.xml", "<core/>")
        z.writestr("docProps/app.xml", "<app/>")
        z.writestr("customXml/item3.xml", "<custom>FINGERPRINT_FOR_TEST</custom>")
        z.writestr("customXml/itemProps3.xml", "<props>METADATA_LEAK</props>")
        z.writestr("customXml/item17.xml", "<custom>another</custom>")
    raw = buf.getvalue()
    assert b"FINGERPRINT_FOR_TEST" in raw

    clean, _ = scrub_office(raw)
    with zipfile.ZipFile(io.BytesIO(clean)) as z:
        names = z.namelist()
        all_content = b"".join(z.read(n) for n in names)
    # Parts still present (manifest integrity)
    assert "customXml/item3.xml" in names
    assert "customXml/itemProps3.xml" in names
    assert "customXml/item17.xml" in names
    # But content is gone
    assert b"FINGERPRINT_FOR_TEST" not in all_content
    assert b"METADATA_LEAK" not in all_content


def test_styles_xml_left_untouched() -> None:
    """styles.xml should NOT be transformed (rsids there are tool fingerprints
    that can break custom-style inheritance if stripped). Confirm w:rsid
    attrs inside styles.xml survive scrubbing."""
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED) as z:
        z.writestr("[Content_Types].xml", '<?xml version="1.0"?><Types xmlns="x"/>')
        z.writestr("docProps/core.xml", "<core/>")
        z.writestr("docProps/app.xml", "<app/>")
        z.writestr(
            "word/styles.xml",
            '<?xml version="1.0"?>'
            '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:style w:rsid="00ABC123" w:styleId="MyStyle"/>'
            '</w:styles>',
        )
    raw = buf.getvalue()
    clean, _ = scrub_office(raw)
    with zipfile.ZipFile(io.BytesIO(clean)) as z:
        styles = z.read("word/styles.xml")
    # rsid in styles.xml deliberately preserved
    assert b"00ABC123" in styles
